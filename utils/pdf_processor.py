import os
from pypdf import PdfReader
from docx import Document
from langchain_core.documents import Document as LCDocument
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import FAISS
from langchain_community.embeddings import HuggingFaceEmbeddings

# Define where we want to save our vector database locally
VECTOR_STORE_DIR = "vectorstore/faiss_index"

def extract_text_from_file(file_path: str) -> str:
    """Reads PDF, DOCX, or TXT files and returns their raw text contents."""
    ext = os.path.splitext(file_path)[1].lower()
    text = ""
    
    if ext == ".pdf":
        reader = PdfReader(file_path)
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
                
    elif ext == ".docx":
        doc = Document(file_path)
        for paragraph in doc.paragraphs:
            if paragraph.text:
                text += paragraph.text + "\n"
                
    elif ext == ".txt":
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            text = f.read()
            
    return text.strip()

def process_and_store_document(file_path: str, doc_type: str):
    """Extracts text, splits it into semantic chunks, generates embeddings, and saves to FAISS."""
    # Step 1: Extract raw string data from the document
    raw_text = extract_text_from_file(file_path)
    if not raw_text:
        return False
    
    # Step 2: Wrap the text into a standard LangChain Document object with metadata
    # Metadata helps us filter searches between "incidents" vs "rca" documents later
    filename = os.path.basename(file_path)
    doc = LCDocument(page_content=raw_text, metadata={"source": filename, "type": doc_type})
    
    # Step 3: Split text into small chunks so embeddings stay accurate and fit inside AI context windows
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=1000, 
        chunk_overlap=200
    )
    chunks = text_splitter.split_documents([doc])
    
    # Step 4: Load our free local embedding model
    embeddings = HuggingFaceEmbeddings(model_name="all-MiniLM-L6-v2")
    
    # Step 5: Check if a FAISS database already exists. If yes, add to it; if no, create a new one.
    if os.path.exists(VECTOR_STORE_DIR):
        db = FAISS.load_local(VECTOR_STORE_DIR, embeddings, allow_dangerous_deserialization=True)
        db.add_documents(chunks)
        db.save_local(VECTOR_STORE_DIR)
    else:
        db = FAISS.from_documents(chunks, embeddings)
        db.save_local(VECTOR_STORE_DIR)
        
    return True